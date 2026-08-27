from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "/Users/sudlabha/Desktop/paw/outputs/serena_revenue_model_answers.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def set_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D0D7DE")
        borders.append(tag)
    tbl_pr.append(borders)


def style_run(run, size=9.2, bold=False, color="111827"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.15)
section.bottom_margin = Cm(1.05)
section.left_margin = Cm(1.05)
section.right_margin = Cm(1.05)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.04

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("PAWJAI Revenue Model - Answers for Serena")
style_run(title_run, size=15, bold=True, color="0B2545")

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(5)
lead_run = lead.add_run(
    "Core answer: insurance should be the main monetisation model because it is driven by completed adoptions and policy conversion, not daily active users. Ads can support discovery and retention, but they are a volume game: in the base model one active user is worth only about ฿1.40/month from ads/commerce, while insurance has a blended customer price of ฿264/month."
)
style_run(lead_run, size=10.2)

rows = [
    (
        "Question Serena raised",
        "Short answer with figure",
        "What we should say",
    ),
    (
        "How many active users are needed for ads?",
        "At Year 5 cost, ads need about 140,574 MAU to break even. The model reaches 101,349 MAU by Year 5, so ads alone are not the safest core model.",
        "Ads need scale, stickiness and direct pet-commerce. They should be a secondary layer, not the main investor story.",
    ),
    (
        "How do we retain users after adoption?",
        "Base assumption: 60% monthly retention and 0.55 viral coefficient. Retention features: My Adopted Pets, vaccine/insurance reminders, donations, document wallet and care profile.",
        "The retention logic should also support insurance later, so features are not wasted if we move from ads to conversion.",
    ),
    (
        "Which pet insurance partners and pricing?",
        "Best first partners: Pet Friend/Thai Vivat, then Rabbit Care x Falcon. Current market is roughly ฿50-฿1,250/month; PAWJAI default should be ฿199/month, with blended price ฿264/month.",
        "Phase 1: partner/broker model. Assumption: 18% broker commission, PAWJAI keeps 50% = 9% of premium. Phase 2: own broker licence captures 18%. Phase 3: own product/MGA captures more but adds claims risk.",
    ),
    (
        "Why is PAWJAI better than insurers selling alone?",
        "Insurers acquire through SEO, broker pages, discounts, LINE/call centres and generic ads. PAWJAI reaches users at the exact adoption moment.",
        "We know the dog profile: age, size, vaccine status, medical notes and special needs. That lets us recommend the right plan and give insurers warmer, lower-CAC leads.",
    ),
    (
        "What is the insurance financial projection?",
        "Insurance starts Month 13. Year 2: 162 new policies, 146 active policies, ฿19k PAWJAI revenue. Year 3: 496 new policies, 553 active policies, ฿198k revenue. Year 5: 1,436 new policies, 2,237 active policies, ฿5.76m revenue, ฿385k EBITDA.",
        "The key driver is adoptions x attach rate. Base case: 80% underwriter acceptance x 70% policy purchase = 56% net attach rate.",
    ),
    (
        "What socio-economic segment are we targeting?",
        "Not only low-income adopters. The viable insurance buyer is middle to upper-middle income: experienced pet owners, same-sex couples and child-free couples treating pets as family.",
        "Use ฿99/month as an entry plan, but make ฿199/month the default recommendation and ฿399/month the upgrade.",
    ),
    (
        "Are there adjacent high-ticket products?",
        "Yes: vet-care memberships, adoption starter kits, shelter SaaS/management and pet-commerce bundles. A 1,000-policy book at ฿264/month equals ฿264k/month gross premium.",
        "Model C should be adjacent to care, not random ads: products that monetize the adopted-dog lifecycle are aligned with insurance.",
    ),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_borders(table)
widths = [4.55, 6.45, 7.9]
hdr = table.rows[0].cells
for i, text in enumerate(rows[0]):
    p = hdr[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, size=8.8, bold=True, color="0B2545")
    set_cell_shading(hdr[i], "E8EEF5")
    set_cell_width(hdr[i], widths[i])
    hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

for row in rows[1:]:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        style_run(r, size=8.35, bold=(i == 0), color="111827")
        set_cell_width(cells[i], widths[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

close = doc.add_paragraph()
close.paragraph_format.space_before = Pt(5)
close.paragraph_format.space_after = Pt(0)
r = close.add_run(
    "Meeting conclusion: present PAWJAI as an adoption-to-care conversion platform. The three-year goal is to prove partner insurance conversion and shelter supply; the five-year upside comes from owning more of the insurance value chain."
)
style_run(r, size=10.0, bold=True, color="0B2545")

source = doc.add_paragraph()
source.paragraph_format.space_after = Pt(0)
r = source.add_run(
    "Sources used: PAWjai Financial Projection v2.xlsx, financial projection CSV, Serena business-model notes PDF/transcript, and Thai pet-insurance competitor research. Raw M4A audio was checked, but not separately transcribed because the PDF already contained the transcript."
)
style_run(r, size=7.7, color="4B5563")

doc.save(OUT)
print(OUT)
